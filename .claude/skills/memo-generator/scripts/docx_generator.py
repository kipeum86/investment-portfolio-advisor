#!/usr/bin/env python3
"""
DOCX Memo Generator — Step 11 of the pipeline.

Generates a Word document investment memo from portfolio-recommendation.json.

Usage:
    python docx_generator.py --recommendation <path> --environment <path> --profile <path> --output <path>
"""

import argparse
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


ASSET_CLASS_LABELS = {
    "us_equity": "US Equity",
    "kr_equity": "KR Equity",
    "bonds": "Bonds",
    "alternatives": "Alternatives",
    "cash": "Cash",
}


def create_allocation_table(allocation: dict) -> list[list[str]]:
    """Convert allocation dict to table rows."""
    rows = []
    for key in ["us_equity", "kr_equity", "bonds", "alternatives", "cash"]:
        label = ASSET_CLASS_LABELS.get(key, key)
        value = allocation.get(key, 0)
        rows.append([label, f"{value:.1f}%"])
    return rows


def create_holdings_table(holdings: list[dict]) -> list[list[str]]:
    """Convert holdings list to table rows."""
    rows = []
    for h in holdings:
        rows.append([
            h.get("ticker", ""),
            h.get("name", ""),
            f"{h.get('weight', 0):.1f}%",
            h.get("rationale", ""),
            h.get("source_tag", ""),
        ])
    return rows


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = str(val)

    return table


def generate_memo(
    recommendation: dict,
    environment: dict,
    user_profile: dict,
    output_path: str,
):
    """Generate the DOCX investment memo."""
    doc = Document()

    # -- Document style defaults --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # -- Cover Page --
    title = doc.add_heading("Investment Portfolio Recommendation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    budget = user_profile.get("budget", 0)
    currency = user_profile.get("budget_currency", "KRW")
    period = user_profile.get("investment_period_months", 0)
    risk = user_profile.get("risk_tolerance", "moderate")

    if currency == "KRW" and budget >= 10000:
        budget_str = f"{budget / 10000:,.0f}만원"
    else:
        budget_str = f"{currency} {budget:,.0f}"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{budget_str} | {period}개월 | {risk} profile")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date: {recommendation.get('recommendation_timestamp', '')[:10]}")

    grade = environment.get("data_grade", "C")
    grade_para = doc.add_paragraph()
    grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grade_para.add_run(f"Data Confidence: Grade {grade}")

    doc.add_page_break()

    # -- 1. Executive Summary --
    doc.add_heading("1. Executive Summary", level=1)
    regime = recommendation.get("regime", environment.get("regime_classification", "unknown"))
    doc.add_paragraph(
        f"This report presents three portfolio options for an investor with "
        f"{budget_str} over {period} months with a {risk} risk tolerance. "
        f"The current market environment is classified as '{regime}'. "
        f"Three options are provided: Aggressive, Moderate, and Conservative, "
        f"each differentiated by equity exposure and risk positioning."
    )

    # -- 2. Market Environment Analysis --
    doc.add_heading("2. Market Environment Analysis", level=1)
    doc.add_paragraph(
        f"Regime Classification: {regime}"
    )
    rationale = environment.get("regime_rationale", "")
    if rationale:
        doc.add_paragraph(f"Rationale: {rationale}")

    # -- 3-5. Portfolio Options --
    for i, opt in enumerate(recommendation.get("options", []), start=3):
        opt_id = opt.get("option_id", f"option_{i}")
        label = opt.get("label", opt_id.title() + " Portfolio")
        doc.add_heading(f"{i}. {label}", level=1)

        # Allocation table
        doc.add_heading("Allocation", level=2)
        alloc_rows = create_allocation_table(opt.get("allocation", {}))
        _add_table(doc, ["Asset Class", "Weight"], alloc_rows)
        doc.add_paragraph()  # spacing

        # Holdings table
        doc.add_heading("Key Holdings", level=2)
        holdings_rows = create_holdings_table(opt.get("holdings", []))
        _add_table(doc, ["Ticker", "Name", "Weight", "Rationale", "Source"], holdings_rows)
        doc.add_paragraph()

        # Scenarios
        doc.add_heading("Scenario Analysis", level=2)
        scenarios = opt.get("scenarios", {})
        scenario_rows = []
        for sc in ["bull", "base", "bear"]:
            s = scenarios.get(sc, {})
            scenario_rows.append([
                sc.title(),
                f"{s.get('expected_return', 0):.1f}%",
                f"{s.get('probability', 0)}%",
                s.get("assumptions", ""),
            ])
        _add_table(doc, ["Scenario", "Expected Return", "Probability", "Assumptions"], scenario_rows)

        rr_score = opt.get("risk_return_score", 0)
        doc.add_paragraph(f"\nRisk/Return Score: {rr_score:.2f}")
        doc.add_paragraph()

    # -- Risk Analysis --
    doc.add_heading(f"{len(recommendation.get('options', [])) + 3}. Risk Analysis", level=1)
    for risk_item in recommendation.get("key_risks", []):
        p = doc.add_paragraph()
        p.add_run(f"Risk: ").bold = True
        p.add_run(risk_item.get("risk", ""))
        doc.add_paragraph(f"Mechanism: {risk_item.get('mechanism', '')}")
        doc.add_paragraph(f"Severity: {risk_item.get('severity', 'medium')}")

    # -- Disclaimer --
    doc.add_heading("Disclaimer", level=1)
    disclaimer = recommendation.get("disclaimer", "This is not investment advice. For informational purposes only.")
    p = doc.add_paragraph(disclaimer)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Generate DOCX investment memo")
    parser.add_argument("--recommendation", required=True)
    parser.add_argument("--environment", required=True)
    parser.add_argument("--profile", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    with open(args.recommendation) as f:
        rec = json.load(f)
    with open(args.environment) as f:
        env = json.load(f)
    with open(args.profile) as f:
        profile = json.load(f)

    generate_memo(rec, env, profile, args.output)
    print(f"Memo generated: {args.output}")


if __name__ == "__main__":
    main()
